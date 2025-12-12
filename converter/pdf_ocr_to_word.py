"""
PDF OCR 转 Word 核心转换模块

支持多种 OCR 引擎和文本提取模式：
- Tesseract OCR
- PaddleOCR
- macOS Vision OCR
- PDF 文本层直读

Author: Excalibur9527
"""

import os
import re
import threading
from concurrent.futures import ThreadPoolExecutor
from typing import Optional, List, Callable, Tuple, Literal

import pdfplumber
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None  # type: ignore

try:
    from ocrmac import ocrmac
except ImportError:  # pragma: no cover
    ocrmac = None  # type: ignore

try:
    # PaddleOCR 作为可选依赖，引入失败时仍允许使用 Tesseract
    from paddleocr import PaddleOCR
except ImportError:  # pragma: no cover - 仅在未安装 PaddleOCR 时触发
    PaddleOCR = None  # type: ignore

_paddle_ocr_instances = {}


def _get_paddle_ocr(lang: str):
    """
    懒加载 + 缓存 PaddleOCR 实例，避免每页都重新加载模型。
    lang 通常使用：
        - 中文：'ch'
        - 英文：'en'
    """
    if PaddleOCR is None:
        raise RuntimeError("未安装 PaddleOCR，请先安装 paddleocr 和 paddlepaddle 再使用 ocr_engine='paddle'")

    if lang not in _paddle_ocr_instances:
        # use_angle_cls=True 识别时考虑文字方向；
        # 旧版本 PaddleOCR 可能不支持 show_log 参数，这里不传，避免 Unknown argument 报错
        _paddle_ocr_instances[lang] = PaddleOCR(lang=lang, use_angle_cls=True)
    return _paddle_ocr_instances[lang]


def _ocr_image_to_text_tesseract(image: Image.Image, lang: str) -> str:
    """
    使用 Tesseract 对单张图片做 OCR。
    """
    config = "--psm 3"
    text = pytesseract.image_to_string(image, lang=lang, config=config)
    return text.strip()


def _ocr_image_to_text_paddle(image: Image.Image, lang: str) -> str:
    """
    使用 PaddleOCR 对单张图片做 OCR。
    """
    ocr = _get_paddle_ocr(lang)
    # result: list[ [ [box, (text, score)], ... ] ]
    result = ocr.ocr(image, cls=True)
    lines: List[str] = []
    for page in result:
        for line in page:
            text, score = line[1]
            # 可以按置信度过滤，例如 score > 0.3
            lines.append(text)
    return "\n".join(lines).strip()


def _clean_chinese_spacing(text: str) -> str:
    """
    将中文之间的多余空格清理掉，例如把“人 设”修复为“人设”。
    """
    return re.sub(r"([\u4e00-\u9fa5])\s+([\u4e00-\u9fa5])", r"\1\2", text)


def _format_page_text(text: str) -> str:
    """
    简单的段落格式化：
    - 去掉首尾空白
    - 合并多余空格
    - 去除中文之间多余空格
    - 按行合并：行末不是句号/问号/感叹号/冒号/引号时，和下一行合并为一句
    """
    text = text.strip()
    if not text:
        return ""

    # 先按行拆分
    lines = [ln.strip() for ln in text.splitlines() if ln.strip() != ""]

    merged: List[str] = []
    buf = ""
    end_punct = ("。", "！", "？", "!", "?", "：", ":", "；", ";", "…", "”", "\"", "）", ")")

    for ln in lines:
        ln = _clean_chinese_spacing(ln)
        # 合并多余空格
        ln = re.sub(r"\s+", " ", ln)

        if not buf:
            buf = ln
            continue

        # 如果上一行以句末标点结束，则换段，否则接在后面
        if buf.endswith(end_punct):
            merged.append(buf)
            buf = ln
        else:
            buf = f"{buf} {ln}"

    if buf:
        merged.append(buf)

    return "\n\n".join(merged)


def _extract_text_layer_pages(
    pdf_path: str,
    remove_tokens: Optional[List[str]] = None,
    progress_callback: Optional[Callable[[int, int], None]] = None,
    auto_format: bool = True,
) -> List[str]:
    """
    使用 pdfplumber 提取 PDF 内置文本层（无 OCR），返回按页的文本列表。
    """
    texts: List[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        if progress_callback is not None:
            progress_callback(0, total)
        processed = 0
        for idx, page in enumerate(pdf.pages, start=1):
            content = page.extract_text()
            if content:
                content = _clean_chinese_spacing(content)
                if remove_tokens:
                    for token in remove_tokens:
                        content = content.replace(token, "")
                if auto_format:
                    content = _format_page_text(content)
                texts.append(content.strip())
            else:
                texts.append("")  # 保留空页占位

            processed += 1
            if progress_callback is not None:
                progress_callback(processed, total)
    return texts


def extract_pdf_text_to_docx(
    pdf_path: str,
    output_path: str,
    remove_tokens: Optional[List[str]] = None,
    progress_callback: Optional[Callable[[int, int], None]] = None,
    auto_format: bool = True,
    font_name: str = "SimSun",
    font_size_pt: int = 12,
) -> str:
    """
    直接提取 PDF 的文本层并写入 Word（不做 OCR）。
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    os.makedirs(output_dir, exist_ok=True)

    page_texts = _extract_text_layer_pages(
        pdf_path=pdf_path,
        remove_tokens=remove_tokens,
        progress_callback=progress_callback,
        auto_format=auto_format,
    )

    doc = Document()
    _apply_default_font(doc, font_name=font_name, font_size_pt=font_size_pt)
    for idx, text in enumerate(page_texts, start=1):
        doc.add_paragraph(text or "")
        if idx != len(page_texts):
            doc.add_page_break()

    if not output_path.lower().endswith(".docx"):
        output_path = output_path + ".docx"
    doc.save(output_path)
    return os.path.abspath(output_path)


def _pdf_to_images(
    pdf_path: str,
    dpi: int = 300,
    poppler_path: Optional[str] = None,
) -> List[Image.Image]:
    """
    将 PDF 每一页转换成 PIL.Image 列表。
    dpi 建议 300，OCR 效果会更好但也更耗时。
    """
    images = convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path)
    return images


def _ocr_pages_in_parallel(
    images: List[Image.Image],
    lang: str,
    ocr_engine: Literal["tesseract", "paddle"] = "tesseract",
    max_workers: Optional[int] = None,
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> List[str]:
    """
    使用线程池对多页图片并行做 OCR，返回按页顺序排列的文本列表。

    progress_callback(current, total):
        - current: 已完成页数
        - total: 总页数
    """
    total = len(images)

    # 结果按 index 存放，确保顺序稳定
    results: List[Optional[str]] = [None] * total

    processed = 0
    lock = threading.Lock()

    # 提前通知总页数，便于外部初始化进度条
    if progress_callback is not None:
        progress_callback(0, total)

    def worker(item: Tuple[int, Image.Image]) -> Tuple[int, str]:
        nonlocal processed
        idx, img = item
        if ocr_engine == "paddle":
            text = _ocr_image_to_text_paddle(img, lang=lang) or ""
        else:
            text = _ocr_image_to_text_tesseract(img, lang=lang) or ""
        if progress_callback is not None:
            with lock:
                processed += 1
                progress_callback(processed, total)
        return idx, text

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        for idx, text in executor.map(worker, enumerate(images)):
            results[idx] = text

    # 类型断言：此时结果列表应无 None
    return [t or "" for t in results]


def convert_pdf_to_docx(
    pdf_path: str,
    output_path: str,
    lang: str = "chi_sim",
    dpi: int = 300,
    poppler_path: Optional[str] = None,
    ocr_engine: Literal["tesseract", "paddle"] = "tesseract",
    max_workers: Optional[int] = None,
    progress_callback: Optional[Callable[[int, int], None]] = None,
    auto_format: bool = True,
    font_name: str = "SimSun",
    font_size_pt: int = 12,
) -> str:
    """
    将 PDF 通过 OCR 转成 Word 文档。

    :param pdf_path: 输入 PDF 文件路径
    :param output_path: 输出的 .docx 文件路径
    :param lang: 语言/模型代码：
        - 当 ocr_engine='tesseract' 时，如 "chi_sim"、"eng"、"chi_sim+eng"
        - 当 ocr_engine='paddle'   时，常用 "ch"（中文）、"en"（英文）
    :param dpi: 渲染 PDF 时的分辨率，越高识别越准但越慢、越占内存
    :param poppler_path: 可选的 poppler 可执行文件所在目录（Windows 常用）
    :param ocr_engine: OCR 引擎类型："tesseract" 或 "paddle"
    :param max_workers: 线程池最大并发数，默认 None 使用系统默认（通常为 CPU 核心数 * 5）
    :param progress_callback: 可选进度回调函数 progress_callback(current, total)
    :return: 实际保存的 .docx 文件路径
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    # 确保输出目录存在
    output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    os.makedirs(output_dir, exist_ok=True)

    # 1. PDF -> 图像列表
    images = _pdf_to_images(pdf_path, dpi=dpi, poppler_path=poppler_path)

    if not images:
        raise ValueError("PDF 转换为图片失败，未得到任何页面。")

    # 2. 对每一页做 OCR（多线程）
    page_texts = _ocr_pages_in_parallel(
        images=images,
        lang=lang,
        ocr_engine=ocr_engine,
        max_workers=max_workers,
        progress_callback=progress_callback,
    )

    if auto_format:
        page_texts = [_format_page_text(text) for text in page_texts]

    # 3. 写入 Word 文档
    doc = Document()
    _apply_default_font(doc, font_name=font_name, font_size_pt=font_size_pt)
    for idx, text in enumerate(page_texts, start=1):
        # 每页一个大段落，可根据需要拆分为多段：for line in text.splitlines()
        doc.add_paragraph(text or "")

        # 不是最后一页时添加分页符
        if idx != len(page_texts):
            doc.add_page_break()

    # 4. 保存 Word 文件
    if not output_path.lower().endswith(".docx"):
        output_path = output_path + ".docx"

    doc.save(output_path)
    return os.path.abspath(output_path)


def convert_pdf_to_docx_mac_vision(
    pdf_path: str,
    output_path: str,
    dpi: int = 300,
    progress_callback: Optional[Callable[[int, int], None]] = None,
    auto_format: bool = True,
    font_name: str = "SimSun",
    font_size_pt: int = 12,
) -> str:
    """
    使用 macOS 原生 Vision OCR（通过 ocrmac）将 PDF 转成 Word。
    仅在 macOS 且安装了 PyMuPDF、ocrmac 的情况下可用。
    """
    if fitz is None:
        raise RuntimeError("未安装 PyMuPDF(fitz)，请先 pip install PyMuPDF")
    if ocrmac is None:
        raise RuntimeError("未安装 ocrmac，请先 pip install ocrmac")

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    os.makedirs(output_dir, exist_ok=True)

    doc_pdf = fitz.open(pdf_path)
    total = len(doc_pdf)
    if progress_callback is not None:
        progress_callback(0, total)

    page_texts: List[str] = []
    processed = 0

    for idx, page in enumerate(doc_pdf, start=1):
        pix = page.get_pixmap(dpi=dpi)
        img_bytes = pix.tobytes("png")
        temp_image_path = f"_tmp_mac_ocr_page_{idx}.png"
        with open(temp_image_path, "wb") as f:
            f.write(img_bytes)

        try:
            ocr = ocrmac.OCR(
                temp_image_path,
                recognition_level="accurate",
                language_preference=["zh-Hans", "en-US"],
            )
            results = ocr.recognize()
            page_text = ""
            if results:
                page_text = "\n".join([item[0] for item in results])
            page_texts.append(page_text)
        finally:
            if temp_image_path and os.path.exists(temp_image_path):
                os.remove(temp_image_path)

        processed += 1
        if progress_callback is not None:
            progress_callback(processed, total)

    docx = Document()
    _apply_default_font(docx, font_name=font_name, font_size_pt=font_size_pt)

    if auto_format:
        page_texts = [_format_page_text(text) for text in page_texts]

    for idx, text in enumerate(page_texts, start=1):
        docx.add_paragraph(text or "")
        if idx != len(page_texts):
            docx.add_page_break()

    if not output_path.lower().endswith(".docx"):
        output_path = output_path + ".docx"
    docx.save(output_path)
    return os.path.abspath(output_path)


def _apply_default_font(doc: Document, font_name: str, font_size_pt: int) -> None:
    """
    设置全局默认字体和字号（同时设置西文字体和中文字体）。
    """
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    # 设置中文字体（东亚）
    r_pr = style._element.rPr
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_name)


__all__ = ["convert_pdf_to_docx", "extract_pdf_text_to_docx", "convert_pdf_to_docx_mac_vision"]


